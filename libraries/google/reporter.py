import matplotlib.pyplot as plt
from docx import Document
from libraries.utilities.utils import get_owned_domains, get_domain_from_email
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING

class Reporter:
    def __init__(self):
        """
        Initializes the Reporter class with default values.
        """
        self.quarantined = 0
        self.delivered = 0
        self.viewed = 0
        self.marked_as_spam = 0
        self.bounced = 0
        self.emails_sent_to_external_domains = 0
        self.total_emails = 0
        self.start_time = None
        self.author = None
        self.author_title = None
        self.author_email = None
        self.author_date = None
        self.mitigations = None
        self.title = None
        self.senders = set()
        self.owned_domains = get_owned_domains()

    def generate_report(self, filename: str, entries: list) -> None:
        """
        Generates a report document after updating counts and creating a pie chart image.

        :param filename: The name of the .docx file where the report will be saved.
        :param entries: A list of email log entries to be processed for the report.
        """
        self.__update_numbers(entries)
        self.__generate_pie_chart()
        self.__generate_docx(filename)

    def ingest_custom_attributes(self, attributes: dict) -> None:
        """
        Loads the custom_attributes dictionary to improve the content of the report.
        :param attributes: A dictionary containing custom attributes for the report.
        """
        for attr, value in attributes.items():
            if value:
                setattr(self, attr, value)

    @property
    def report_title(self) -> str:
        """
        Generates a title for the report, using the custom title if provided.
        :return: The title of the report as a string.
        """
        return self.title if self.title else f"{self.start_time.strftime('%m/%d/%Y')} Phishing Incident"

    @property
    def summary_text(self) -> str:
        """
        Generates a summary of the provided events.
        :return: The summary text of the incident report.
        """
        formatted_start_time = self.start_time.strftime("%I:%M %p")
        sender_text = self.senders.pop() if len(self.senders) == 1 else "multiple users"
        internal_count = self.total_emails - self.emails_sent_to_external_domains

        summary_intro = (
            f"At {formatted_start_time}, an email from {sender_text} was identified as a phishing attempt. "
            f"It was sent to a total of {self.total_emails} recipients, including "
            f"{internal_count} within our managed domains and {self.emails_sent_to_external_domains} external organizations."
            if self.emails_sent_to_external_domains > 0 else
            f"At {formatted_start_time}, an email from {sender_text} was identified as a phishing attempt. "
            f"It was sent to {self.total_emails} recipients within our managed domains."
        )
        viewed_text = f"Of these, {self.viewed} emails were viewed by the recipients." if self.viewed > 0 else \
            "Fortunately, none of the recipients within our managed domains opened the email."
        
        return f"{summary_intro} {viewed_text}"

    def __update_numbers(self, entries: list) -> None:
        """
        Updates email status counts based on entries, processing each unique user only once.
        :param entries: A list of email log entries to be processed.
        """
        processed_users = {}
        external_recipients = set()

        for email_log in entries:
            recipient = email_log.recipient_address
            self.senders.add(email_log.sender)
            self.start_time = min(self.start_time, email_log.start_date) if self.start_time else email_log.start_date

            if recipient not in processed_users:
                domain = get_domain_from_email(recipient)
                if domain not in self.owned_domains:
                    external_recipients.add(recipient)

            if email_log.was_email_viewed:
                processed_users[recipient] = "viewed"
            elif email_log.was_email_bounced:
                processed_users.setdefault(recipient, "bounced")
            elif email_log.was_email_quarantined and processed_users.get(recipient) not in ["viewed", "bounced"]:
                processed_users[recipient] = "quarantined"
            elif email_log.was_email_delivered:
                processed_users.setdefault(recipient, "delivered")

        self.viewed = sum(1 for status in processed_users.values() if status == "viewed")
        self.bounced = sum(1 for status in processed_users.values() if status == "bounced")
        self.quarantined = sum(1 for status in processed_users.values() if status == "quarantined")
        self.delivered = sum(1 for status in processed_users.values() if status == "delivered")
        self.total_emails = len(processed_users)
        self.emails_sent_to_external_domains = len(external_recipients)

    def __generate_docx(self, filename: str) -> None:
        """
        Generates a .docx report containing information about the incident.
        :param filename: The name of the .docx file where the report will be saved.
        """
        doc = Document()
        doc.add_heading(self.report_title, level=1)
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(self.summary_text)

        doc.add_heading("Email Distribution Overview", level=2)
        doc.add_picture("chart.png", width=Inches(4.0))

        if self.mitigations:
            doc.add_heading("Mitigations Taken", level=2)
            style = doc.styles['List Bullet']
            style.paragraph_format.line_spacing = 1.0
            for mitigation in self.mitigations:
                if not mitigation:
                    continue
                paragraph = doc.add_paragraph()
                paragraph.style = style
                paragraph.add_run(mitigation)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)

        if self.author:
            doc.add_heading("Report By:", level=4)
            doc.add_paragraph(self.author)
            if self.author_email:
                doc.add_paragraph(self.author_email)
            if self.author_title:
                doc.add_paragraph(self.author_title)
            if self.author_date:
                doc.add_paragraph(self.author_date)
                
        doc.save(filename)

    def __generate_pie_chart(self) -> None:
        """
        Generates a static pie chart image for the report.
        """
        labels = ['Viewed', 'Bounced', 'Quarantined', 'Delivered']
        values = [self.viewed, self.bounced, self.quarantined, self.delivered]
        filtered_labels = [label for label, value in zip(labels, values) if value > 0]
        filtered_values = [value for value in values if value > 0]

        plt.figure(figsize=(6, 6))
        plt.pie(filtered_values, labels=filtered_labels, autopct=lambda pct: f"{int(round(pct * sum(filtered_values) / 100.0))}", startangle=140)
        plt.title(f"Total Emails Delivered: {self.total_emails}")
        plt.savefig("chart.png")
        plt.close()
